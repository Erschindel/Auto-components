import docx
import re

#create document
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

#format from list of tuples to list. the ORs in the regex result in tuples when put through re.findall
def formatList(raw, clean, size) :
    for group in range(len(raw)) :
        for dex in range(size) :
            if raw[group][dex] != "" :
                clean.append(raw[group][dex])
    return

claims = getText("test2.docx")
components = []
firstComponents = []
lowerComponents = []
firstWithSubs = []
paragraphs = []

#regular expressions; identify and separate components
r1 = "(?<=\na )(.*)(?=;)|(?<=\nan )(.*)(?=;)|(?<=\n)(at.*)(?=;)"
firstCompRaw = re.findall(r1, claims)
formatList(firstCompRaw, firstComponents, 3)
for comp in firstComponents:
    components.append(comp)

r2 = "(?<=\nthe )(.*)(?= further comprises)|(?<=\nthe )(.*)(?= comprises)"
firstWithSubsRaw = re.findall(r2, claims)
formatList(firstWithSubsRaw, firstWithSubs, 2)

#r3 identifies all subcomponent names, including subs of subs
r3 = "(?<=comprises a )(.*?)(?=, )|(?<=comprises a )(.*?)(?= and)|(?<=comprises a )(.*?)(?=;)|(?<=comprises )(at.*?)(?=,)|(?<=comprises )(at.*?)(?= and)|(?<=comprises )(at.*?)(?=;)|(?<= an )(.*?)(?=, )|(?<=comprises an )(.*?)(?= and)|(?<=comprises an )(.*?)(?=;)|(?<= and an )(.*?)(?=;)|(?<=, )(at.*?)(?=,)|(?<=, and )(at.*?)(?=;)|(?<=, a )(.*?)(?=,)|(?<= and a )(.*?)(?=;)|(?<= and )(at.*?)(?=;)"
subCompsRaw = re.findall(r3, claims)
formatList(subCompsRaw, lowerComponents, 15)

#create list of paragraphs
r4 = "(?<=\n)(.*?)(?=\n)"
paragraphs.append(re.findall(r4, claims))

#fill components dictionary with appropriate subcomponents
r5 = "(?<=the )(.*)(?= further comprises)|(?<=the )(.*)(?= comprises)"
r6 = "(?<=the )(.*)(?= further comprises)"

secondComponents = []
thirdComponents = []
fourthComponents = []
fifthComponents = []

for line in range(len(paragraphs[0])) :
    key = re.findall(r5, paragraphs[0][line])
    cleanKey = []
    formatList(key, cleanKey, 2)
    check = re.findall(r6, paragraphs[0][line])
    inlineComps = re.findall(r3, paragraphs[0][line])
    #check if the line contains new subcomponents
    if key :
        chunk = []
        #ordered list of parent components
        parentComponent = cleanKey[len(cleanKey) - 1]

        #find new subcomponents in given line
        for each in range(len(inlineComps)) :
            for comp in range(len(inlineComps[each])) :
                if inlineComps[each][comp] :
                    chunk.append(inlineComps[each][comp])
        for i in range(len(chunk)) :
            #separate components into levels
            if(parentComponent in firstComponents) :
                secondComponents.append(chunk[i])
            elif(parentComponent in secondComponents) :
                thirdComponents.append(chunk[i])
            elif(parentComponent in thirdComponents) :
                fourthComponents.append(chunk[i])
            elif(parentComponent in fourthComponents) :
                fifthComponents.append(chunk[i])

            nextComp = components[components.index(parentComponent)]
        #add all subcomponents to components[] in correct order
            #second-level components
            if(check and (chunk[i] in secondComponents)) :
                components.insert(components.index(firstComponents[firstComponents.index(parentComponent) + 1]), f"{chunk[i]}")
            #third-level components. NOTE: can make function to shorten third- and fourth-level component handling maybe?
            elif(check and (chunk[i] in thirdComponents)) :
                if(parentComponent in secondComponents) :
                    while nextComp not in secondComponents or nextComp not in firstComponents :
                        nextComp = components[components.index(nextComp) + 1]
                        if(components.index(nextComp) == len(components) - 1) :
                            components.append(f"{chunk[len(chunk) - i - 1]}")
                            break
                        elif(nextComp in secondComponents or nextComp in firstComponents) :
                            break
                    if(components.index(nextComp) != len(components) - 1) :
                        components.insert(components.index(nextComp), f"{chunk[i]}")
                elif(parentComponent in firstComponents) :
                    if(nextComp not in firstComponents) :
                        if(components.index(nextComp) != len(components) - 1) :
                            while(nextComp not in secondComponents or nextComp not in firstComponents) :
                                nextComp = components[components.index(parentComponent) + 1]
                        else :
                            components.append(f"{chunk[len(chunk) - i - 1]}")
                    if(components.index(nextComp) != len(components) - 1) :
                        components.insert(components.index(nextComp) - 1, f"{chunk[len(chunk) - i - 1]}")
                else :
                    components.append(f"{chunk[len(chunk) - i - 1]}")
            #fourth-level components
            elif(check and (chunk[i] in fourthComponents)) :
                if(parentComponent in thirdComponents) :
                    while nextComp not in thirdComponents or nextComp not in secondComponents or nextComp not in firstComponents :
                        nextComp = components[components.index(nextComp) + 1]
                        if(components.index(nextComp) == len(components) - 1) :
                            components.append(f"{chunk[len(chunk) - i - 1]}")
                            break
                        elif(nextComp in thirdComponents or nextComp in secondComponents or nextComp in firstComponents) :
                            break
                    if(components.index(nextComp) != len(components) - 1) :
                        components.insert(components.index(nextComp), f"{chunk[i]}")
                elif(parentComponent in secondComponents) :
                    while nextComp not in secondComponents or nextComp not in firstComponents :
                        nextComp = components[components.index(nextComp) + 1]
                        if(components.index(nextComp) == len(components) - 1) :
                            components.append(f"{chunk[len(chunk) - i - 1]}")
                            break
                        elif(nextComp in secondComponents or nextComp in firstComponents) :
                            break
                    if(components.index(nextComp) != len(components) - 1) :
                        components.insert(components.index(nextComp), f"{chunk[i]}")
                elif(parentComponent in firstComponents) :
                    if(nextComp not in firstComponents) :
                        if(components.index(nextComp) != len(components) - 1) :
                            while(nextComp not in secondComponents or nextComp not in firstComponents) :
                                nextComp = components[components.index(parentComponent) + 1]
                        else :
                            components.append(f"{chunk[len(chunk) - i - 1]}")
                    if(components.index(nextComp) != len(components) - 1) :
                        components.insert(components.index(nextComp) - 1, f"{chunk[len(chunk) - i - 1]}")
                else :
                    components.append(f"{chunk[len(chunk) - i - 1]}")
            else:
                components.insert(components.index(parentComponent) + 1, f"{chunk[len(chunk) - i - 1]}")

# print(f"\n{components}")

#generate components word document with component numbers
componentsDoc = docx.Document()
for component in range(0, len(components)) :
    if(components[component] in firstComponents) :
        componentsDoc.add_paragraph(f"{components[component]} {component + 1}")
    elif(components[component] in secondComponents) :
        componentsDoc.add_paragraph(f"\t{components[component]} {component + 1}")
    elif(components[component] in thirdComponents) :
        componentsDoc.add_paragraph(f"\t\t{components[component]} {component + 1}")
    elif(components[component] in fourthComponents) :
        componentsDoc.add_paragraph(f"\t\t\t{components[component]} {component + 1}")
    elif(components[component] in fifthComponents) :
        componentsDoc.add_paragraph(f"\t\t\t\t{components[component]} {component + 1}")
    else :
        break
componentsDoc.save("components.docx")
