import docx
from makeList import components as components
from makeList import firstComponents as firstComponents
from makeList import secondComponents as secondComponents
from makeList import thirdComponents as thirdComponents
from makeList import fourthComponents as fourthComponents
from makeList import fifthComponents as fifthComponents
#generate components word document with component numbers
componentsDoc = docx.Document()
for component in range(0, len(components)) :
    if(components[component] in firstComponents) :
        componentsDoc.add_paragraph(f"{components[component]} {component + 1}")
    elif(components[component] in secondComponents) :
        componentsDoc.add_paragraph(f"\t{components[component]} {component + 1}")
    elif(components[component] in thirdComponents) :
        componentsDoc.add_paragraph(f"\t\t{components[component]} {component + 1}")
    elif(components[component] in fourthComponents) :
        componentsDoc.add_paragraph(f"\t\t\t{components[component]} {component + 1}")
    elif(components[component] in fifthComponents) :
        componentsDoc.add_paragraph(f"\t\t\t\t{components[component]} {component + 1}")
    else :
        print("error: sixth-level components present")
        break
componentsDoc.save("../components.docx")
